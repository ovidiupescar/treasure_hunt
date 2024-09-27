import qrcode
from docx import Document
from docx.shared import Inches
from io import BytesIO
import pandas as pd
from PIL import Image
import random

# Sample list of names and links
names_and_links = [
    ("John Doe", "https://example.com/johndoe"),
    ("Jane Smith", "https://example.com/janesmith"),
    ("John Doe", "https://example.com/johndoe"),
    ("Jane Smith", "https://example.com/janesmith"),
    ("John Doe", "https://example.com/johndoe"),
    ("Jane Smith", "https://example.com/janesmith"),
    ("John Doe", "https://example.com/johndoe"),
    ("Jane Smith", "https://example.com/janesmith"),
    ("John Doe", "https://example.com/johndoe"),
    ("Jane Smith", "https://example.com/janesmith"),
    ("John Doe", "https://example.com/johndoe"),
    ("Jane Smith", "https://example.com/janesmith"),
    ("John Doe", "https://example.com/johndoe"),
    ("Jane Smith", "https://example.com/janesmith"),
    # Add more names and links as needed
]

def generate_qr_code(link):
    """Generate a QR code image for the given link."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(link)
    qr.make(fit=True)
    img = qr.make_image(fill='black', back_color='white')

    # Convert QR code image to a Pillow image for saving
    img_pil = img.convert("RGB")

    # Save the image to a BytesIO object
    img_byte_arr = BytesIO()
    img_pil.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)

    return img_byte_arr

def create_docx_with_qr_codes(names_and_links, docx_filename):
    """Create a .docx file containing names with their corresponding QR codes."""
    doc = Document()

    doc.add_heading("Titlu")

    for name, link in names_and_links:
        # Add the name as a new paragraph
        doc.add_paragraph(name)

        # Generate QR code for the link
        qr_code_image = generate_qr_code(link)

        # Save the QR code to a temporary file and insert it in the docx
        qr_code_img = Image.open(qr_code_image)

        # Save the image temporarily in BytesIO for inserting into docx
        qr_code_image_temp = BytesIO()
        qr_code_img.save(qr_code_image_temp, format="PNG")
        qr_code_image_temp.seek(0)

        # Insert the QR code into the document
        doc.add_picture(qr_code_image_temp, width=Inches(1.5))

        # Add a line break after each name-QR pair
        doc.add_paragraph()

    # Save the docx file
    doc.save(docx_filename)

def main():
    # Generate the .docx file
    create_docx_with_qr_codes(names_and_links, "names_with_qr_codes.docx")

def gen_link(n, statie):
    return f"https://oradea.artorius.uk/question/?group_number={n}&question_number={statie}"

def get_links():
    df = pd.read_pickle("intrebari.pkl")
    
    result_licu = []
    result_explo = []
    licu_list = []
    explo_list = []

    df.index = pd.RangeIndex(start=1, stop=len(df) + 1, step=1)
    
    for i, row in df.iterrows():
        explo_list.append((i, row['Statie'], row['Grupa']))
        if row['Grupa'] == 'licu':
            licu_list.append((i, row['Statie'], row['Grupa']))
    
    for i, current_tuple in enumerate(licu_list):
        # Copy the original list and remove the current tuple from the list
        temp_list = licu_list[:i] + licu_list[i+1:]

        # Shuffle the remaining tuples
        random.shuffle(temp_list)

        # Insert the current tuple at the first position
        new_list = [current_tuple] + temp_list
        
        # Append this new list to the result_list
        result_licu.append(new_list)

    for i, current_tuple in enumerate(explo_list):
        (_, _, group_type) = current_tuple
        if group_type == 'explo':
            # Copy the original list and remove the current tuple from the list
            temp_list = explo_list[:i] + explo_list[i+1:]

            # Shuffle the remaining tuples
            random.shuffle(temp_list)

            # Insert the current tuple at the first position
            new_list = [current_tuple] + temp_list
            
            # Append this new list to the result_list
            result_explo.append(new_list)

    return result_licu, result_explo

print(get_links())
#main()