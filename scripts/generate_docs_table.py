import qrcode
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from PIL import Image
import pandas as pd
import random

def get_lists():
    df = pd.read_pickle("intrebari.pkl")
    
    result_licu = []
    result_explo = []
    licu_list = []
    explo_list = []

    df.index = pd.RangeIndex(start=1, stop=len(df) + 1, step=1)
    
    for i, row in df.iterrows():
        explo_list.append((i, row['Statie'], row['Grupa']))
        if row['Grupa'] == 'licu':
            licu_list.append((i, row['Statie'], row['Grupa']))
    
    for i, current_tuple in enumerate(licu_list):
        # Copy the original list and remove the current tuple from the list
        temp_list = licu_list[:i] + licu_list[i+1:]

        # Shuffle the remaining tuples
        random.shuffle(temp_list)

        # Insert the current tuple at the first position
        new_list = [current_tuple] + temp_list
        
        # Append this new list to the result_list
        result_licu.append(new_list)

    for i, current_tuple in enumerate(explo_list):
        (_, _, group_type) = current_tuple
        if group_type == 'explo':
            # Copy the original list and remove the current tuple from the list
            temp_list = explo_list[:i] + explo_list[i+1:]

            # Shuffle the remaining tuples
            random.shuffle(temp_list)

            # Insert the current tuple at the first position
            new_list = [current_tuple] + temp_list
            
            # Append this new list to the result_list
            result_explo.append(new_list)

    random.shuffle(result_licu)
    random.shuffle(result_explo)

    return result_licu, result_explo


def set_table_borders_none(table):
    """Set the table borders to have no borders."""
    tbl = table._tbl  # Access the underlying XML element
    tblPr = tbl.tblPr  # Get the table properties

    # If tblPr is None, create it and append it to tbl
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    # Remove existing borders if any
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)

    # Create new tblBorders element with no borders
    tblBorders = OxmlElement('w:tblBorders')

    border_sides = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']

    for side in border_sides:
        border_el = OxmlElement(f'w:{side}')
        border_el.set(qn('w:val'), 'nil')  # Set border to 'nil' (no border)
        tblBorders.append(border_el)

    tblPr.append(tblBorders)

# Function to set cell borders
def set_cell_borders(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#00FF00", "space": "0"},
        left={"sz": 24, "val": "dashed", "color": "0000FF", "space": "0"},
        right={"sz": 12, "val": "dashed", "color": "00FF00", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in edge_data:
                element.set(qn(f'w:{key}'), str(edge_data[key]))

def generate_qr_code(link):
    """Generate a QR code image for the given link."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(link)
    qr.make(fit=True)
    img = qr.make_image(fill='black', back_color='white')

    # Convert QR code image to a Pillow image for saving
    img_pil = img.convert("RGB")

    # Save the image to a BytesIO object
    img_byte_arr = BytesIO()
    img_pil.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)

    return img_byte_arr

def set_cell_vertical_alignment(cell, align="center"):
    """Set vertical alignment of a table cell."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(qn('w:val'), align)
    tcPr.append(tcVAlign)

def set_cell_white_border(cell):
    """Set the cell border to white."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')  # White color
        borders.append(border)

    tcPr.append(borders)

def add_header(doc, name):
    """Add a header to the document with the given document name."""
    # Access the section where we will add the header
    section = doc.sections[0]
    
    # Get the header for this section
    header = section.header
    
    # Add a paragraph in the header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = name
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center-align the header


def add_page_number(doc):
    """Add page numbers to the footer."""
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Page "
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add field for page numbers
    run = footer_paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_docx_with_custom_layout(group_number, names_and_links, docx_filename):
    """Create a .docx file containing a title, custom margins, and a table with 4 columns (Name, QR Code, Name, QR Code)."""
    doc = Document()

    # Adjust margins (make them smaller)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Add title to the document
    if group_number <= 17:
        group_type = 'Licu'
    else:
        group_type = 'Explo'
    title = doc.add_heading(f'Grupa {group_number} ______________', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add page numbers to the footer
    add_page_number(doc)
    add_header(doc, f"{group_type} {group_number}")

    # Add a table with 4 columns: Name, QR Code, Name, QR Code
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'

    # Set the table borders to "No Borders"
    set_table_borders_none(table)

    for i in range(0, len(names_and_links), 2):
        # Add a new row in the table
        row_cells = table.add_row().cells

        # First pair: Name and QR Code
        name1, link1 = names_and_links[i]
        row_cells[0].text = name1
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the name
        set_cell_vertical_alignment(row_cells[0], 'center')  # Vertically center the name

        qr_code_image1 = generate_qr_code(link1)
        qr_code_img1 = Image.open(qr_code_image1)
        qr_code_image_temp1 = BytesIO()
        qr_code_img1.save(qr_code_image_temp1, format="PNG")
        qr_code_image_temp1.seek(0)
        qr_code_paragraph1 = row_cells[1].paragraphs[0]
        run1 = qr_code_paragraph1.add_run()
        run1.add_picture(qr_code_image_temp1, width=Inches(1.5))  # Increase QR code size
        qr_code_paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the QR code

        #set_cell_white_border(row_cells[0])  # White border between name and QR code
        #set_cell_white_border(row_cells[1])

        # Second pair (if available): Name and QR Code
        if i + 1 < len(names_and_links):
            name2, link2 = names_and_links[i + 1]
            row_cells[2].text = name2
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the name
            set_cell_vertical_alignment(row_cells[2], 'center')  # Vertically center the name

            qr_code_image2 = generate_qr_code(link2)
            qr_code_img2 = Image.open(qr_code_image2)
            qr_code_image_temp2 = BytesIO()
            qr_code_img2.save(qr_code_image_temp2, format="PNG")
            qr_code_image_temp2.seek(0)
            qr_code_paragraph2 = row_cells[3].paragraphs[0]
            run2 = qr_code_paragraph2.add_run()
            run2.add_picture(qr_code_image_temp2, width=Inches(1.5))  # Increase QR code size
            qr_code_paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the QR code

            #set_cell_white_border(row_cells[2])  # White border between name and QR code
            #set_cell_white_border(row_cells[3])

        # Add extra space between rows (add a paragraph with custom spacing)
        extra_space = doc.add_paragraph()
        extra_space_format = extra_space.paragraph_format
        extra_space_format.space_before = Pt(12)  # Set custom space between rows
    
    # Now set borders only on the first Name and QR Code cells
    first_row_cells = table.rows[0].cells
    set_cell_borders(first_row_cells[0], top={"val": "single", "sz": "6", "color": "000000"},
                     bottom={"val": "single", "sz": "6", "color": "000000"},
                     left={"val": "single", "sz": "6", "color": "000000"},
                     right={"val": "nil", "sz": "0", "color": "000000"})
    set_cell_borders(first_row_cells[1], top={"val": "single", "sz": "6", "color": "000000"},
                     bottom={"val": "single", "sz": "6", "color": "000000"},
                     left={"val": "nil", "sz": "0", "color": "000000"},
                     right={"val": "single", "sz": "6", "color": "000000"})

    # Save the document
    doc.save(docx_filename)

# Generate the .docx file without row restriction per page
#create_docx_with_custom_layout(names_and_links, "names_with_qr_codes_custom_no_row_limit.docx")

def generate_link(nr_statie, nr_grupa):
    return f"https://oradea.artorius.uk/question/?group_number={nr_grupa}&question_number={nr_statie}"

def main():
    licu_statii, explo_statii = get_lists()
    #create_docx_with_custom_layout(group_number, names_and_links, "names_with_qr_codes_custom_no_row_limit.docx")

    for i, fisa in enumerate (licu_statii):
        nr_grupa = i+1
        print(f"Generam pt grupa {nr_grupa}")
        names_and_links = []
        for statie in fisa:
            nr_statie, nume_statie, grupa = statie
            names_and_links.append((nume_statie, generate_link(nr_statie, nr_grupa)))

        create_docx_with_custom_layout(nr_grupa, names_and_links, f"generate/L{nr_grupa}.docx")
    
    for i, fisa in enumerate (explo_statii):
        nr_grupa = i + 17 + 1
        print(f"Generam pt grupa {nr_grupa}")
        names_and_links = []
        for statie in fisa:
            nr_statie, nume_statie, grupa = statie
            names_and_links.append((nume_statie, generate_link(nr_statie, nr_grupa)))

        create_docx_with_custom_layout(nr_grupa, names_and_links, f"generate/E{nr_grupa}.docx")
        
main()